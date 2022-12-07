def fill(name, Email, arr):
    import docx
    import os
    doc = docx.Document("reportTemplate.docx")
    hospital="FCAI HOSPITAL"
    for row in doc.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text == "Patient’s Name":
                    paragraph.text = "Patient’s Name : " + name
                elif paragraph.text == "Email":
                    paragraph.text = "Email : " + Email
                elif paragraph.text == "Hospital":
                    paragraph.text = "Hospital : " + hospital


    for paragraph in doc.paragraphs:
        if paragraph.text == "   -   When were they admitted to your hospital?":
            paragraph.text = "   -   When were they admitted to your hospital?\n\n" + arr[0]
        elif paragraph.text == " -  Reason for admission and medical diagnosis":
            paragraph.text = " -  Reason for admission and medical diagnosis\n\n" + arr[1]
        elif paragraph.text == "-  Past medical history (if known)":
            paragraph.text = "-  Past medical history (if known)\n\n" + arr[2]
        elif paragraph.text == "- Progress on ward":
            paragraph.text = "- Progress on ward\n\n" + arr[3]
        elif paragraph.text == "- Current clinical condition":
            paragraph.text = "- Current clinical condition\n\n" + arr[4]
        elif paragraph.text == " - Prognosis and prospects for rehabilitation":
            paragraph.text = " - Prognosis and prospects for rehabilitation\n\n" + arr[5]
        elif paragraph.text == "-  Relevant laboratory results, x-rays etc":
            paragraph.text = "-  Relevant laboratory results, x-rays etc\n\n" + arr[6]
        elif paragraph.text == "- Current medication":
            paragraph.text = "- Current medication\n\n" + arr[7]
        elif paragraph.text == "- Arrangements to follow up":
            paragraph.text = "- Arrangements to follow up\n\n" + arr[8]

    doc.save("{}.docx".format(name))
    path = os.path.abspath(os.getcwd())
    print(path)
    os.startfile(path + "/" +name+".docx")